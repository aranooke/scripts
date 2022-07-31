import math,pikepdf;
import docx;
import aspose.words as aw
from pygoogletranslation import Translator
from pdfminer.high_level import extract_text
from docx import Document
from docx.shared import Pt
from docx2pdf import convert



document = Document()
translator = Translator()
text = extract_text("Akunin.pdf",page_numbers = range(100));
arr = text.split(".");
for i in arr:
    try:
        print("Programm is working");
        document.add_paragraph(i);
        document.add_paragraph(translator.translate(i,dest = 'ru').text);
        # print(i);
        # print(translator.translate(i,dest = 'ru').text);
    except Exception:
        pass;
print('Programm complete');
t = translator.translate('Hello world',dest ='ru');
document.save("t.docx");
"""
doc = aw.Document("txt.docx");
doc.save("text.pdf");"""
